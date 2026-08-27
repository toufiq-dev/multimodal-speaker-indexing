#!/usr/bin/env python3
"""Build thesis DOCX + PDF matching the Daffodil International University sample.

Usage:  .venv/bin/python build_thesis_docx.py
Output: ~/Developer/multimodal-speaker-indexing-Thesis-Report/thesis/Thesis_Report.docx
        ~/Developer/multimodal-speaker-indexing-Thesis-Report/thesis/Thesis_Report.pdf
"""
from __future__ import annotations

import re
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Paths ──────────────────────────────────────────────────────────────
THESIS_DIR = (
    Path.home()
    / "Developer"
    / "multimodal-speaker-indexing-Thesis-Report"
    / "thesis"
)
FIG_DIR = THESIS_DIR / "figures"
OUT_DOCX = THESIS_DIR / "Thesis_Report.docx"
OUT_PDF = THESIS_DIR / "Thesis_Report.pdf"

# ── Formatting constants ───────────────────────────────────────────────
FONT = "Times New Roman"
BODY = Pt(12)
LINE_SP = 1.5
MARGIN = {"top": Inches(1), "bottom": Inches(1), "left": Inches(1.5), "right": Inches(1)}
HDR_TEXT = "©Daffodil International University"

# Figure map: section number -> (filename, caption)
FIGURE_MAP = {
    "3.1": ("fig3_1_pipeline.png", "Figure 3.1: End-to-end multimodal pipeline architecture"),
    "3.2": ("fig3_2_cascade.png", "Figure 3.2: Identity resolution cascade (P0–P5)"),
    "3.3": ("fig3_3_midpoint.png", "Figure 3.3: Word-midpoint containment vs IoU assignment"),
    "4.1": ("fig4_1_duplicate.png", "Figure 4.1: Duplicate text rate before and after fix"),
    "4.2": ("fig4_2_cuelen.png", "Figure 4.2: Cue length distribution (pre-fix vs post-fix)"),
    "4.3": ("fig4_3_confidence.png", "Figure 4.3: Confidence distribution — untrained gate vs deterministic cascade"),
    "4.4": ("fig4_4_ablation.png", "Figure 4.4: Ablation — cpWER by modality (lower is better)"),
    "4.5": ("fig4_5_dashboard.png", "Figure 4.5: Fusion health dashboard (Global TV 45m)"),
    "4.6": ("fig4_6_yolo_selection.png", "Figure 4.6: YOLO model selection on DAWN (mAP comparison)"),
}

# ── Helpers ────────────────────────────────────────────────────────────

def _shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _hdr_footer(section, roman=False):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run(HDR_TEXT + "    ")
    r.font.name = FONT
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    hp._p.append(OxmlElement("w:tab"))
    # Page number field
    rn = hp.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    rn._r.append(fc)
    ri = hp.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE \\* ROMAN " if roman else " PAGE "
    ri._r.append(it)
    re_ = hp.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    re_._r.append(fe)


def _pf(p, align=None, sb=0, sa=0, indent=None):
    f = p.paragraph_format
    if align is not None:
        f.alignment = align
    f.space_before = Pt(sb)
    f.space_after = Pt(sa)
    f.line_spacing = LINE_SP
    if indent is not None:
        f.first_line_indent = indent


def _run(p, text, bold=False, italic=False, size=BODY, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = size
    if color:
        r.font.color.rgb = color
    return r


def _clean(text: str) -> str:
    """Strip markdown artifacts."""
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    t = re.sub(r"`(.+?)`", r"\1", t)
    t = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", t)
    if re.match(r"^\s*[-*_]{3,}\s*$", t):
        return ""
    return t.strip()


def _add_table(doc, rows):
    """Add a real Word table from rows (list of list of str). First row is header."""
    if not rows or len(rows) < 2:
        # Fallback: at least header + 1 row needed for a table; otherwise render as body
        for r in rows:
            _body(doc, "  |  ".join(r), indent=False)
        return
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        c = hdr_cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shading(c, "1a3c6e")
        c.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Body rows
    for row in body:
        cells = table.add_row().cells
        for i, txt in enumerate(row):
            if i >= len(cells):
                continue
            c = cells[i]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(9)
            if i == 0:
                _shading(c, "eaf2f8")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_figure(doc, section_number, width_inches=6.0):
    """Embed figure for a given section number if mapping exists."""
    entry = FIGURE_MAP.get(section_number)
    if not entry:
        return
    fname, caption = entry
    fpath = FIG_DIR / fname
    if not fpath.exists():
        p = doc.add_paragraph()
        _pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=3)
        _run(p, f"[Missing figure: {fname}]", italic=True, size=Pt(9), color=RGBColor(0xC0, 0x39, 0x2B))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(fpath), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


# ── Content loaders ────────────────────────────────────────────────────

def _load_front_matter() -> dict[str, list[str]]:
    """Parse 00_front_matter.md into {heading: [lines]}."""
    text = (THESIS_DIR / "00_front_matter.md").read_text("utf-8")
    secs: dict[str, list[str]] = {}
    key: str | None = None
    lines: list[str] = []
    for line in text.split("\n"):
        if re.match(r"^# [^#]", line):
            continue  # skip document title
        m = re.match(r"^##\s+(.+)", line)
        if m:
            if key is not None:
                secs[key] = lines
            key = m.group(1).strip()
            lines = []
            continue
        m2 = re.match(r"^###\s+(.+)", line)
        if m2:
            lines.append(f"**{m2.group(1)}**")
            continue
        if key is not None:
            lines.append(line)
    if key is not None:
        secs[key] = lines
    return secs


def _load_chapter(filename: str) -> list[dict]:
    """Parse chapter markdown into [{number, title, lines}]."""
    text = (THESIS_DIR / filename).read_text("utf-8")
    secs = []
    cur_num = None
    cur_title = None
    cur_lines: list[tuple[str, str]] = []
    for line in text.split("\n"):
        if line.startswith("# CHAPTER"):
            continue
        m = re.match(r"^## (\d+\.\d+)\s+(.*)", line)
        if m:
            if cur_num is not None:
                secs.append({"number": cur_num, "title": cur_title, "lines": cur_lines})
            cur_num, cur_title = m.group(1), m.group(2)
            cur_lines = []
            continue
        m2 = re.match(r"^###\s+(.+)", line)
        if m2:
            cur_lines.append(("h3", m2.group(1)))
            continue
        cur_lines.append(("body", line))
    if cur_num is not None:
        secs.append({"number": cur_num, "title": cur_title, "lines": cur_lines})
    return secs


def _load_refs() -> list[str]:
    text = (THESIS_DIR / "references.md").read_text("utf-8")
    return [_clean(l) for l in text.split("\n") if l.strip() and not l.startswith("# ")]

# ── Document builders ──────────────────────────────────────────────────

def _make_doc() -> Document:
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = BODY
    s.paragraph_format.line_spacing = LINE_SP
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = MARGIN["top"]
    sec.bottom_margin = MARGIN["bottom"]
    sec.left_margin = MARGIN["left"]
    sec.right_margin = MARGIN["right"]
    sec.header_distance = Inches(0.5)
    sec.footer_distance = Inches(0.5)
    # Core properties
    core = doc.core_properties
    core.title = "Multimodal Speaker Indexing in Bengali Talk-show Videos Using Audio-Visual Fusion"
    core.author = "Toufiqur Rahman (252-25-013)"
    core.subject = "MSc Thesis — Daffodil International University"
    core.keywords = "speaker indexing, multimodal fusion, Bengali NLP, diarization, face recognition"
    core.created = datetime(2026, 8, 27, tzinfo=timezone.utc)
    core.modified = datetime.now(timezone.utc)
    core.revision = 2
    return doc


def _title_page(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    for p in sec.first_page_header.paragraphs:
        p.clear()
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    _run(p, "Multimodal Speaker Indexing in Bengali Talk-show Videos\nUsing Audio-Visual Fusion",
         bold=True, size=Pt(18))
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=24)
    _run(p, "BY", bold=True, size=Pt(14))
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=6)
    _run(p, "Toufiqur Rahman", size=Pt(14))
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=48)
    _run(p, "ID: 252-25-013")
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=24)
    _run(p, "This Report is Presented in Partial Fulfillment of the Requirements for\n"
            "The Degree of Masters of Science in Computer Science and Engineering")
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=6)
    _run(p, "Supervised By", bold=True)
    for _ in range(2):
        doc.add_paragraph()
    for txt, b in [("Professor Dr. Sheak Rashed Haider Noori", True), ("Professor & Head", False),
                    ("Department of CSE", False), ("Daffodil International University", False)]:
        p = doc.add_paragraph()
        _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=3)
        _run(p, txt, bold=b)
    for _ in range(3):
        doc.add_paragraph()
    for txt, b in [("DAFFODIL INTERNATIONAL UNIVERSITY", True),
                    ("DHAKA, BANGLADESH", False),                    ("AUGUST 2026", False)]:
        p = doc.add_paragraph()
        _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=3)
        _run(p, txt, bold=b, size=Pt(14) if b else Pt(12))
    doc.add_page_break()


def _fm_page(doc, title, lines, roman=True):
    """Front-matter page."""
    _hdr_footer(doc.sections[-1], roman=roman)
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sb=24, sa=18)
    _run(p, title, bold=True, size=Pt(14))
    for raw in lines:
        line = _clean(raw.rstrip())
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        _pf(p, sa=6, indent=Inches(0.5))
        _run(p, line)


def _chapter(doc, num, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sb=48, sa=24)
    _run(p, f"CHAPTER {num}", bold=True, size=Pt(16))
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sa=24)
    _run(p, title.upper(), bold=True, size=Pt(14))


def _section(doc, number, title):
    p = doc.add_paragraph()
    _pf(p, sb=18, sa=12)
    _run(p, f"{number} {title}", bold=True)


def _body(doc, text, indent=True):
    line = _clean(text)
    if not line:
        return
    p = doc.add_paragraph()
    _pf(p, sa=6, indent=Inches(0.5) if indent else None)
    _run(p, line)


# ── Main build ─────────────────────────────────────────────────────────

FM_ORDER = [
    "Approval", "Declaration", "Acknowledgement", "Abstract",
    "Table of Contents", "List of Figures", "List of Tables",
]

CHAPTERS = [
    (1, "Introduction", "chapter1_introduction.md"),
    (2, "Background", "chapter2_background.md"),
    (3, "Research Methodology", "chapter3_methodology.md"),
    (4, "Experimental Results and Discussion", "chapter4_results_and_discussion.md"),
    (5, "Impact on Society, Environment and Sustainability", "chapter5_impact.md"),
    (6, "Conclusion and Future Work", "chapter6_conclusion.md"),
]


def build():
    doc = _make_doc()
    _title_page(doc)

    # Front matter
    fm = _load_front_matter()
    fm_lower = {k.lower(): v for k, v in fm.items()}
    for title in FM_ORDER:
        tl = title.lower()
        content = None
        for hdr, lines in fm_lower.items():
            if tl in hdr or hdr in tl:
                content = lines
                break
        _fm_page(doc, title, content or [f"[{title}]"])

    # Switch to arabic numbering for chapters
    _hdr_footer(doc.sections[-1], roman=False)

    # Chapters
    for num, title, fname in CHAPTERS:
        secs = _load_chapter(fname)
        _chapter(doc, num, title)
        for sd in secs:
            _section(doc, sd["number"], sd["title"])
            # Buffer for consecutive markdown tables
            table_buffer: list[list[str]] = []
            def flush_table():
                if table_buffer:
                    _add_table(doc, table_buffer)
                    table_buffer.clear()
            for ltype, ltxt in sd["lines"]:
                if ltype == "h3":
                    flush_table()
                    p = doc.add_paragraph()
                    _pf(p, sb=12, sa=6)
                    _run(p, ltxt, bold=True, italic=True)
                else:
                    stripped = ltxt.strip()
                    if not stripped:
                        flush_table()
                        doc.add_paragraph()
                    elif stripped.startswith("```"):
                        flush_table()
                        pass  # code fence
                    elif stripped.startswith("- ") or stripped.startswith("* "):
                        flush_table()
                        p = doc.add_paragraph()
                        _pf(p, sa=3, indent=Inches(0.25))
                        _run(p, "• " + _clean(stripped[2:]))
                    elif re.match(r"^\d+\.\s", stripped):
                        flush_table()
                        p = doc.add_paragraph()
                        _pf(p, sa=3, indent=Inches(0.25))
                        _run(p, _clean(stripped))
                    elif stripped.startswith("|"):
                        cells = [c.strip() for c in stripped.split("|") if c.strip()]
                        if all(re.match(r"^[-:]+$", c) for c in cells):
                            continue  # separator row
                        table_buffer.append(cells)
                    else:
                        flush_table()
                        _body(doc, stripped)
            flush_table()
            # Insert figure for this section if mapped
            _add_figure(doc, sd["number"])

    # References
    doc.add_page_break()
    p = doc.add_paragraph()
    _pf(p, WD_ALIGN_PARAGRAPH.CENTER, sb=48, sa=24)
    _run(p, "REFERENCES", bold=True, size=Pt(16))
    for ref in _load_refs():
        if not ref:
            continue
        p = doc.add_paragraph()
        _pf(p, sa=6, indent=Inches(0.5))
        _run(p, ref)

    doc.save(str(OUT_DOCX))
    print(f"✓ DOCX  {OUT_DOCX}  ({OUT_DOCX.stat().st_size // 1024} KB)")
    # Report figure embedding stats
    from docx.opc.constants import CONTENT_TYPE as CT
    n_images = sum(1 for rel in doc.part.rels.values() if rel.target_ref and "image" in rel.target_ref)
    # Alternative count via inline shapes
    try:
        n_inline = len(doc.inline_shapes)
    except Exception:
        n_inline = n_images
    n_tables = len(doc.tables)
    print(f"  Tables: {n_tables}  Images: {n_inline}")


def to_pdf():
    try:
        r = subprocess.run(
            ["pandoc", str(OUT_DOCX), "-o", str(OUT_PDF), "--pdf-engine=weasyprint"],
            capture_output=True, text=True, timeout=120,
        )
        if r.returncode == 0:
            print(f"✓ PDF   {OUT_PDF}  ({OUT_PDF.stat().st_size // 1024} KB)")
            return
        print(f"  pandoc stderr: {r.stderr[:500]}")
    except Exception as e:
        print(f"  pandoc: {e}")
    print("  → Open the .docx in Word/LibreOffice and export to PDF manually.")


if __name__ == "__main__":
    build()
    to_pdf()
